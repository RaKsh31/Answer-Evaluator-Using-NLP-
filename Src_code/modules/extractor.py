"""
modules/extractor.py — Text Extraction Module
==============================================
Extracts plain text from various file formats:
  PDF  → pdfplumber (with OCR fallback via Tesseract)
  Image → pytesseract OCR
  DOCX  → python-docx
  TXT   → direct read
  XLSX/XLS → openpyxl / pandas (concatenate all cell values)
"""

import os
import pytesseract
from PIL import Image
import pdfplumber
import docx
import pandas as pd


# ── Dispatcher ────────────────────────────────────────────────────────────────

def extract_text(filepath: str) -> str:
    """
    Route file to the correct extractor based on extension.
    Returns extracted text as a single string.
    """
    ext = os.path.splitext(filepath)[1].lower()

    extractors = {
        ".pdf":  _extract_pdf,
        ".png":  _extract_image,
        ".jpg":  _extract_image,
        ".jpeg": _extract_image,
        ".docx": _extract_docx,
        ".txt":  _extract_txt,
        ".xlsx": _extract_excel,
        ".xls":  _extract_excel,
    }

    extractor = extractors.get(ext)
    if not extractor:
        raise ValueError(f"Unsupported file format: {ext}")

    return extractor(filepath)


# ── PDF Extraction ────────────────────────────────────────────────────────────

def _extract_pdf(filepath: str) -> str:
    """
    Extract text from PDF using pdfplumber.
    Falls back to Tesseract OCR if a page has no selectable text
    (i.e., scanned/image-based PDF pages).
    """
    text_parts = []

    with pdfplumber.open(filepath) as pdf:
        for page_num, page in enumerate(pdf.pages, start=1):
            page_text = page.extract_text()

            if page_text and page_text.strip():
                # Selectable text found — use it directly
                text_parts.append(page_text)
            else:
                # No selectable text — OCR the page image
                page_image = page.to_image(resolution=300).original
                ocr_text = pytesseract.image_to_string(page_image, lang="eng")
                text_parts.append(ocr_text)

    return "\n".join(text_parts)


# ── Image Extraction ──────────────────────────────────────────────────────────

def _extract_image(filepath: str) -> str:
    """
    Run Tesseract OCR on an image file.
    Preprocessing (grayscale) improves accuracy on low-contrast images.
    """
    image = Image.open(filepath).convert("L")  # Convert to grayscale
    text = pytesseract.image_to_string(image, lang="eng")
    return text


# ── DOCX Extraction ───────────────────────────────────────────────────────────

def _extract_docx(filepath: str) -> str:
    """
    Extract text from a .docx Word document paragraph by paragraph.
    Table cells are also captured.
    """
    doc = docx.Document(filepath)
    parts = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text)

    # Extract table cells
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                parts.append(row_text)

    return "\n".join(parts)


# ── TXT Extraction ────────────────────────────────────────────────────────────

def _extract_txt(filepath: str) -> str:
    """Read a plain text file with UTF-8 encoding (fallback to latin-1)."""
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(filepath, "r", encoding="latin-1") as f:
            return f.read()


# ── Excel Extraction ──────────────────────────────────────────────────────────

def _extract_excel(filepath: str) -> str:
    """
    Read all sheets from an Excel file and concatenate all cell values.
    Useful when answer sheets are stored in tabular format.
    """
    all_text = []
    xls = pd.ExcelFile(filepath)

    for sheet_name in xls.sheet_names:
        df = xls.parse(sheet_name, header=None)
        # Convert each row to a space-separated string, skip NaN
        for _, row in df.iterrows():
            row_text = " ".join(str(v) for v in row if pd.notna(v) and str(v).strip())
            if row_text:
                all_text.append(row_text)

    return "\n".join(all_text)
